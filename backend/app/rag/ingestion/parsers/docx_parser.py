"""DOCX 解析器 — 基于 python-docx。"""

from __future__ import annotations

from pathlib import Path

from ..models import ParsedDocument
from .base import BaseParser, register_parser


@register_parser
class DocxParser(BaseParser):
    SUPPORTED_EXTENSIONS = {".docx"}
    NAME = "docx"

    def parse(self, file_path: Path) -> list[ParsedDocument]:
        from docx import Document

        doc = Document(str(file_path))
        paragraphs: list[str] = []
        sections: list[dict] = []
        char_offset = 0

        for para in doc.paragraphs:
            text = para.text
            style = para.style.name if para.style else ""
            if style.startswith("Heading"):
                try:
                    level = int(style.replace("Heading ", ""))
                except ValueError:
                    level = 1
                sections.append({
                    "heading": text,
                    "level": level,
                    "char_start": char_offset,
                })
            paragraphs.append(text)
            char_offset += len(text) + 1

        full_text = "\n".join(paragraphs).strip()

        # 提取表格
        tables: list[dict] = []
        for i, table in enumerate(doc.tables):
            table_text = "\n".join(
                " | ".join(cell.text for cell in row.cells)
                for row in table.rows
            )
            tables.append({
                "table_text": table_text,
                "row_count": len(table.rows),
                "index": i,
            })

        return [ParsedDocument(
            text=full_text,
            metadata={
                "source_path": str(file_path),
                "file_type": "docx",
                "sections": sections,
                "tables": tables,
                "title": file_path.stem,
            },
        )]
